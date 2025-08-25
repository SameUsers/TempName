import os
import re
import pdfplumber
import pandas as pd
from openpyxl import load_workbook
from openpyxl.styles import PatternFill
from PIL import Image, ImageEnhance
import pytesseract
from decimal import Decimal, ROUND_HALF_UP
from docx import Document
from difflib import get_close_matches
import tempfile
import openai


class PdfWorker:
    """Обработчик PDF-файлов для извлечения табличных данных и конвертации в XLSX формат"""

    def __init__(self):
        pass

    def pdf_to_xlsx(self, pdf_path: str, xlsx_path: str, filter_spec: bool = False,
                    remove_edges: bool = False, invoice_lines: bool = False) -> bool:
        with pdfplumber.open(pdf_path) as pdf:
            if invoice_lines:
                return self._process_invoice_pdf(pdf, xlsx_path)
            else:
                return self._process_standard_pdf(pdf, xlsx_path, filter_spec, remove_edges)

    def _process_invoice_pdf(self, pdf, xlsx_path):
        pattern = re.compile(
            r'^(\d+)\s+(\d+)\s+(.*?)\s+(\d+)\s+'
            r'(Stk\.?\s*/\s*\d+\s*\w*)\s+'
            r'([\d,.]+(?:\s*/?\s*[\d,.]*)?)\s+'
            r'([\d,.]+)$'
        )

        all_lines = []
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                all_lines.extend([ln.strip() for ln in text.split("\n") if ln.strip()])

        skip_patterns = [
            r'^Rechnung', r'^Kundennummer', r'^Rechnungsdatum', r'^Lieferdatum',
            r'^IBAN', r'^BIC', r'^UST-Id', r'^Zwischensumme', r'^Vortrag',
            r'^Seite', r'^Amtsgericht', r'^HRB', r'^GF:', r'^Gläubiger-ID',
            r'^www\.', r'^http', r'^Eissing GmbH',
        ]

        allowed_patterns = [
            r'^EAN', r'^Art', r'^Hersteller'
        ]

        customs_pattern = re.compile(r'^Zolltarif-Nr\.?:\s*(\d+)')
        tables_invoice = []

        i = 0
        while i < len(all_lines):
            line = all_lines[i]
            match = pattern.match(line)
            if match:
                groups = list(match.groups())
                description = groups[2]
                customs_code = None
                j = i + 1
                while j < len(all_lines):
                    next_line = all_lines[j]
                    if re.match(r'^\d+\s+\d+\s+', next_line):
                        break
                    if any(re.match(pat, next_line) for pat in skip_patterns):
                        j += 1
                        continue
                    customs_match = customs_pattern.match(next_line)
                    if customs_match:
                        customs_code = customs_match.group(1)
                        j += 1
                        continue
                    if any(re.match(pat, next_line) for pat in allowed_patterns):
                        description += " " + next_line
                        j += 1
                        continue
                    break
                groups[2] = description.strip()
                df = pd.DataFrame([groups + [customs_code]], columns=[
                    "№", "Code", "Description", "Quantity",
                    "Unit/Volume", "PricePerUnit", "TotalPrice",
                    "CustomsCode"
                ])
                tables_invoice.append(df)
                i = j
            else:
                i += 1

        if tables_invoice:
            result_df = pd.concat(tables_invoice, axis=0, ignore_index=True, sort=False)
            result_df.to_excel(xlsx_path, index=False)
            return True
        return False

    def _process_standard_pdf(self, pdf, xlsx_path, filter_spec, remove_edges):
        tables_standard = []
        for page in pdf.pages:
            tables = page.extract_tables()
            for table in tables:
                if not table or len(table) < 2:
                    continue
                df = pd.DataFrame(table)
                if filter_spec:
                    df[0] = pd.to_numeric(df[0], errors='coerce')
                    max_val = df[0].max()
                    df = df[(df[0] >= 1) & (df[0] <= max_val)].reset_index(drop=True)
                if remove_edges and len(df) > 2:
                    df = df.iloc[1:-1].reset_index(drop=True)
                tables_standard.append(df)

        if tables_standard:
            result_df = pd.concat(tables_standard, axis=0, ignore_index=True, sort=False)
            result_df.to_excel(xlsx_path, index=False)
            return True
        return False


class FileGenerate:
    """Генератор итоговых XLSX-файлов на основе шаблонов и данных из различных источников"""

    def __init__(self):
        pass
    
    @staticmethod
    def round_half_up_int(value):
            return int(Decimal(str(value)).quantize(Decimal("1"), rounding=ROUND_HALF_UP))
    

    def fill_invoice(self, template_filename: str, invoice_filename: str, ref_filename: str,
                 pl_filename: str, spec_filename: str, output_filename: str,
                 photos_folder: str = "PHOTOS") -> bool:
        template_path = os.path.join("examples", template_filename)
        invoice_path = os.path.join("xlsx_files", invoice_filename)
        ref_path = os.path.join("examples", ref_filename)
        pl_path = os.path.join("xlsx_files", pl_filename)
        spec_path = os.path.join("xlsx_files", spec_filename)
        output_path = os.path.join("examples", output_filename)

        invoice_df = pd.read_excel(invoice_path)
        ref_df = pd.read_excel(ref_path, header=None)
        pl_df = pd.read_excel(pl_path)
        spec_df = pd.read_excel(spec_path)

        ref_map = dict(zip(ref_df[3], ref_df[2]))

        def safe_column(df, col_index, numeric=False, replace_kan=False):
            values = df.iloc[:, col_index]
            if replace_kan:
                values = values.astype(str).str.replace("Kan", "").str.strip()
            if numeric:
                return [
                    float(str(x).replace(",", ".").strip()) if str(x).strip() not in ["", "nan", "None"] else 0
                    for x in values
                ]
            return values.fillna("").tolist()
        
        # Сравниваем количество строк ДО обработки
        spec_col_B_raw = spec_df.iloc[:, 1].dropna().tolist() 
        invoice_col_C_raw = invoice_df.iloc[:, 2].dropna().tolist()
        spec_count = len(spec_col_B_raw)
        invoice_count = len(invoice_col_C_raw)
        
        if spec_count != invoice_count:
            print(f"ВНИМАНИЕ: Несоответствие количества строк!")
            print(f"Specification_sell.xlsx (колонка B): {spec_count} строк")
            print(f"Invoice_purchase.xlsx (колонка C): {invoice_count} строк")
            
            min_count = min(spec_count, invoice_count)
            print(f"Будет обработано: {min_count} строк")
        else:
            min_count = spec_count
            print(f"Всё в порядке. Списки равны. Значение {min_count}")

        # ОБРЕЗАЕМ данные до min_count чтобы избежать ошибок
        pl_col_B = safe_column(pl_df, 1, replace_kan=True)[:min_count]
        pl_col_E = safe_column(pl_df, 4, numeric=True)[:min_count]
        pl_col_D = safe_column(pl_df, 3, numeric=True)[:min_count]
        pl_col_H = safe_column(pl_df, 7, numeric=True)[:min_count]
        pl_col_I = safe_column(pl_df, 8, numeric=True)[:min_count]

        spec_values = [
            float(re.sub(r'[^0-9,]', '', str(x)).replace(',', '.')) if re.search(r'\d', str(x)) else 0
            for x in spec_df.iloc[:, 4]
        ][:min_count]

        invoice_col_H = safe_column(invoice_df, 7)[:min_count]
        spec_col_B = spec_df.iloc[:, 1].tolist()[:min_count]  # Теперь это обрезанный список

        wb = load_workbook(template_path)
        ws = wb.active
        row_start = 18

        yellow_fill = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")

        # Ограничиваем цикл min_count
        for idx in range(min_count):
            code = invoice_df["CustomsCode"].fillna("").tolist()[idx] if idx < len(invoice_df) else ""
            value_c = ref_map.get(code, "")
            row_excel = row_start + idx

            if not value_c:
                val = str(spec_col_B[idx]) if idx < len(spec_col_B) else ""
                text_c = re.sub(r"[^А-Яа-яЁё]", " ", val)
                text_c = re.sub(r"\s+", " ", text_c).strip()
                value_c = text_c

                for col in range(1, ws.max_column + 1):
                    ws.cell(row=row_excel, column=col).fill = yellow_fill

            ws.cell(row=row_excel, column=3, value=value_c)

        # Все остальные циклы также ограничиваем min_count
        for idx in range(min_count):
            val = spec_df.iloc[:, 1].tolist()[idx] if idx < len(spec_df) else ""
            text = str(val) if val is not None else ""
            text = text.replace("\n", " ").replace("\r", " ")
            text = re.sub(r"[А-Яа-яЁё]", "", text)
            text = re.sub(r"\s+", " ", text).strip()
            ws.cell(row=row_start + idx, column=4, value=text)

        for idx in range(min_count):
            val = pl_col_E[idx] if idx < len(pl_col_E) else 0
            ws.cell(row=row_start + idx, column=7, value=val)

        for idx in range(min_count):
            val = pl_col_D[idx] if idx < len(pl_col_D) else 0
            ws.cell(row=row_start + idx, column=8, value=val)

        for idx in range(min_count):
            val = pl_col_H[idx] if idx < len(pl_col_H) else 0
            ws.cell(row=row_start + idx, column=10, value=self.round_half_up_int(val))

        for idx in range(min_count):
            val = pl_col_I[idx] if idx < len(pl_col_I) else 0
            ws.cell(row=row_start + idx, column=11, value=self.round_half_up_int(val))

        for idx in range(min_count):
            val = spec_values[idx] if idx < len(spec_values) else 0
            ws.cell(row=row_start + idx, column=12, value=val)

        for idx in range(min_count):
            val = invoice_col_H[idx] if idx < len(invoice_col_H) else ""
            ws.cell(row=row_start + idx, column=14, value=val)

        for idx in range(min_count):
            product_name = pl_col_B[idx] if idx < len(pl_col_B) else ""
            try:
                candidate_folders = os.listdir(photos_folder)
            except FileNotFoundError:
                candidate_folders = []
            best_match = get_close_matches(product_name, candidate_folders, n=1, cutoff=0.5)
            if not best_match:
                continue
            folder_path = os.path.join(photos_folder, best_match[0])
            if not os.path.isdir(folder_path):
                continue
            jpg_files = [f for f in os.listdir(folder_path) if f.lower().endswith((".jpg",'.png','.jpeg','.JPEG','.JPG'))][:2]
            made_in_value = "EU"
            for jpg in jpg_files:
                image_path = os.path.join(folder_path, jpg)
                try:
                    img = Image.open(image_path)
                    img = img.resize((img.width * 2, img.height * 2))
                    enhancer = ImageEnhance.Contrast(img)
                    img = enhancer.enhance(3.0)
                    text = pytesseract.image_to_string(img)
                    match = re.search(r"made\s*in\s*([A-Za-z\s]+)", text, re.IGNORECASE)
                    if match:
                        country = match.group(1).strip()
                        made_in_value = country if country.upper() != "EU" else "EU"
                        break
                except Exception:
                    continue
            ws.cell(row=row_start + idx, column=5, value=made_in_value)

        def _to_float(x):
            if x is None:
                return None
            if isinstance(x, (int, float)):
                return float(x)
            s = str(x).strip()
            if s == "" or s.lower() in ("nan", "none"):
                return None
            s = s.replace(" ", "").replace(",", ".")
            try:
                return float(s)
            except Exception:
                return None

        # Также ограничиваем этот цикл
        for idx in range(min_count):
            r = row_start + idx
            g_val = _to_float(ws.cell(row=r, column=7).value)  # G
            h_val = _to_float(ws.cell(row=r, column=8).value)  # H
            if g_val is None and h_val is None:
                continue
            total = (g_val or 0.0) * (h_val or 0.0)
            total = self.round_half_up_int(total) 
            ws.cell(row=r, column=6, value=total)  # F

        wb.save(output_path)
        return True

class DocxFiller:
    """Заполнитель DOCX-документов данными из Excel с поддержкой изображений и веб-поиска ссылок"""

    def __init__(self, openai_api_key="Token"):
        self.api_key = openai_api_key

    def get_official_link(self, product_name: str) -> str:
        try:
            client = openai.OpenAI(api_key=self.api_key)
            prompt = f"Предоставь одну единственную официальную ссылку на продукт с названием: {product_name}. Отправляй только ссылку без дополнительных уточнений и прочего текста."

            response = client.responses.create(
                model="o4-mini",
                input=prompt,
                tools=[{
                    "type": "web_search_preview",
                    "search_context_size": "low",
                    "user_location": {
                        "type": "approximate",
                        "country": "GB",
                        "city": "London",
                        "region": "London"
                    }
                }],
            )

            link_text = response.output_text.strip()
            match = re.search(r"https?://\S+", link_text)
            print(match.group(0))
            if match:
                return match.group(0)
            return ""

        except Exception as e:
            print(f"Не удалось получить ссылку для '{product_name}': {e}")
            return ""


    def _get_product_name_before_comma(self, full_name: str) -> str:
        """Извлекает часть названия продукта до запятой"""
        if ',' in full_name:
            full_name=full_name.replace("TotalEnergies","").replace("Dose","").strip()
            return full_name.split(',')[0].strip()
        return full_name.strip()

    def _find_exact_folder_match(self, product_name: str, folder_list: list) -> str:
        count=0
        """Ищет точное совпадение названия продукта (до запятой) с папками"""
        target_name = self._get_product_name_before_comma(product_name)
        
        # Сначала ищем точное совпадение
        for folder in folder_list:
            folder_name = self._get_product_name_before_comma(folder)
            if folder_name.lower() == target_name.lower():
                print(f'Папка-----{folder_name}')
                print(f'Сравнение------{target_name}')
                count+=1
                print(f'-------------------------------------------------------------------')
                return folder
        
        # Если точного совпадения нет, используем приблизительный поиск

        best_matches = get_close_matches(target_name.replace("TotalEnergies","").replace("Dose","").strip(), folder_list, n=1, cutoff=0.3)
        print(f'Мягкий поиск(что ищем){target_name}')
        print(f'Мягкий поиск (что нашли){best_matches[0]}')
        print(f'-------------------------------------------------------------------')
        return best_matches[0] if best_matches else None

    def fill_table_from_excel(self, template_path: str, excel_path: str,
                              output_path: str, photos_folder: str = "PHOTOS",
                              table_index: int = 0) -> bool:
        if not os.path.exists(template_path):
            print(f"Шаблон Word {template_path} не найден")
            return False
        if not os.path.exists(excel_path):
            print(f"Excel-файл {excel_path} не найден")
            return False

        wb = load_workbook(excel_path, data_only=True)
        ws = wb.active

        values_Start = [str(ws[f"C{row}"].value or "") for row in range(18, 48)]
        values_D = [str(ws[f"D{row}"].value or "") for row in range(18, 48)]
        values_E = [str(ws[f"E{row}"].value or "") for row in range(18, 48)]
        values_F = [str(ws[f"F{row}"].value or "") for row in range(18, 48)]
        values_N = [str(ws[f"N{row}"].value or "") for row in range(18, 48)]

        doc = Document(template_path)
        try:
            table = doc.tables[table_index]
        except IndexError:
            print(f"Таблица с индексом {table_index} не найдена в шаблоне Word")
            return False

        # Получаем список всех папок один раз
        folder_candidates = []
        if os.path.exists(photos_folder):
            folder_candidates = [f for f in os.listdir(photos_folder) 
                               if os.path.isdir(os.path.join(photos_folder, f))]

        for i in range(len(values_D)):
            row_index = i + 1
            if row_index >= len(table.rows):
                print(f"Недостаточно строк в таблице для строки {row_index+1}")
                continue

            link = self.get_official_link(values_D[i])
            cell = table.rows[row_index].cells[1]

            if link:
                cell.text = f"{values_Start[i]}\n{link}"
            else:
                cell.text = f"{values_Start[i]}\n{values_D[i]}"

            table.rows[row_index].cells[2].text = values_D[i]
            table.rows[row_index].cells[3].text = values_E[i]
            table.rows[row_index].cells[4].text = values_F[i]
            table.rows[row_index].cells[5].text = values_N[i]

            # Поиск папки с точным совпадением названия до запятой
            best_match = self._find_exact_folder_match(values_D[i], folder_candidates)
            if best_match:
                folder_path = os.path.join(photos_folder, best_match)
                jpg_files = [f for f in os.listdir(folder_path) if f.lower().endswith((".jpg",'.png','.jpeg','.JPEG','.JPG'))][:2]
                if jpg_files:
                    from docx.shared import Cm
                    images = [Image.open(os.path.join(folder_path, f)) for f in jpg_files]
                    heights = [im.height for im in images]
                    max_h = max(heights)
                    total_w = sum(im.width for im in images)
                    merged = Image.new("RGB", (total_w, max_h), (255, 255, 255))
                    x_offset = 0
                    for im in images:
                        merged.paste(im, (x_offset, 0))
                        x_offset += im.width
                    tmp_path = tempfile.mktemp(suffix=".png")
                    merged.save(tmp_path)
                    cell = table.rows[row_index].cells[6]
                    cell.text = ""
                    run = cell.paragraphs[0].add_run()
                    run.add_picture(tmp_path, width=Cm(7.3), height=Cm(4.95))
                    continue

            table.rows[row_index].cells[6].text = values_N[i]

        doc.save(output_path)
        return True


if __name__ == '__main__':
    file_custom = FileGenerate()
    file_custom.fill_invoice(
        template_filename="123456789_invoice_sell.xlsx",
        invoice_filename='Invoice_purchase.xlsx',
        ref_filename="Справочник.xlsx",
        pl_filename="PL.xlsx",
        spec_filename="Specification_sell.xlsx",
        output_filename='Result.xlsx'
    )